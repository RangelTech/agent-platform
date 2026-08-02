"""QA de memória, documentos, artefatos e isolamento entre empresas.

Cada verificação aqui existe porque o modo de falhar é silencioso: o agente
responde bem, o cliente acredita, e o problema só aparece semanas depois.

- **Artefato**: dizer "gráfico gerado" é fácil; o teste baixa o artefato e
  confere que existe conteúdo. Já foi visto o agente anunciar um gráfico que
  nunca chegou ao chat.
- **Documento**: o agente pode responder de cabeça em vez de ler o arquivo. Os
  documentos deste teste têm um dado inventado e único, que só pode vir do
  arquivo.
- **Isolamento**: o teste tenta, de propósito, ver memória e artefato da outra
  empresa — estando logado na segunda.

Uso:
    LICITA_FIXTURES=<pasta> python scripts/qa_memoria_artefatos.py
"""

import json
import os
import sys
import time
import uuid
from pathlib import Path

import httpx

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

BACKEND = os.environ.get("REGRESSAO_BACKEND", "https://teste-ia-backend-x27vtpiida-uc.a.run.app")
EMAIL = os.environ.get("LICITA_EMAIL", "dono@licita.com")
SENHA = os.environ.get("LICITA_SENHA", "licita-senha-forte-123")
MASTER_EMAIL = os.environ.get("HOMOLOG_MASTER_EMAIL", "master@example.com")
MASTER_SENHA = os.environ.get("HOMOLOG_MASTER_PASSWORD", "admin123")
TEMPLATE = "Inteligência de Licitações"

# Dado inventado: não existe em lugar nenhum além do arquivo. Se o agente
# acertar, leu; se errar, respondeu de cabeça.
SEGREDO_TXT = "O código interno do contrato guarda-chuva da LicitaEnterprisse é LKT-9931-ZX."
SEGREDO_DOCX = "A meta de faturamento da LicitaEnterprisse para 2027 é R$ 48.750.000,00."

resultados: list[dict] = []


def check(grupo: str, nome: str, ok: bool, detalhe: str = "") -> bool:
    resultados.append({"grupo": grupo, "nome": nome, "ok": bool(ok), "detalhe": detalhe[:400]})
    print(f"  [{'ok  ' if ok else 'FALHA'}] {grupo}: {nome}")
    if not ok and detalhe:
        print(f"         {detalhe[:300]}")
    return bool(ok)


def auth(token: str) -> dict:
    return {"Authorization": f"Bearer {token}"}


def _consumir(response) -> dict:
    reply, tools, artifacts, agents, erro, chat_id = "", [], [], [], None, None
    evento = ""
    for linha in response.iter_lines():
        if linha.startswith("event: "):
            evento = linha[7:].strip()
        elif linha.startswith("data: "):
            try:
                dados = json.loads(linha[6:])
            except ValueError:
                continue
            if evento == "chat":
                chat_id = dados.get("chat_id", chat_id)
            elif evento == "tool":
                tools.append(dados)
            elif evento == "artifact":
                artifacts.append(dados)
            elif evento == "agent":
                agents.append(dados)
            elif evento == "done":
                reply = dados.get("text", "")
            elif evento == "error":
                erro = dados.get("detail") or dados
    return {
        "reply": reply,
        "tools": tools,
        "artifacts": artifacts,
        "agents": agents,
        "error": erro,
        "chat_id": chat_id,
    }


def enviar(client, token, mensagem, template_id, chat_id=None) -> dict:
    corpo = {"message": mensagem, "template_id": template_id}
    if chat_id:
        corpo["chat_id"] = chat_id
    with client.stream(
        "POST", "/api/chat/send", json=corpo, headers=auth(token), timeout=600.0
    ) as resposta:
        return _consumir(resposta)


# --------------------------------------------------------------------------


def suite_documentos(client, token, template_id, pasta: Path):
    """Um .txt e um .docx viram conhecimento do agente."""
    print("\n-- documentos: .txt e .docx viram conhecimento do agente --")

    txt = pasta / "contrato_guarda_chuva.txt"
    txt.write_text(
        "Contrato guarda-chuva — LicitaEnterprisse\n\n"
        f"{SEGREDO_TXT}\n\n"
        "Vigência: 2026 a 2028. Responsável: diretoria comercial.\n",
        encoding="utf-8",
    )

    docx = pasta / "metas_2027.docx"
    try:
        from docx import Document

        documento = Document()
        documento.add_heading("Metas LicitaEnterprisse 2027", level=1)
        documento.add_paragraph(SEGREDO_DOCX)
        documento.add_paragraph("Documento interno, uso restrito.")
        documento.save(docx)
    except ImportError:
        check("documentos", "python-docx disponível para gerar o .docx", False)
        return

    enviados = {}
    for caminho in (txt, docx):
        with caminho.open("rb") as fh:
            resposta = client.post(
                "/api/files",
                files={"file": (caminho.name, fh, "application/octet-stream")},
                headers=auth(token),
                timeout=300.0,
            )
        ok = resposta.status_code == 201
        check("documentos", f"upload de {caminho.suffix} aceito", ok, resposta.text)
        if ok:
            enviados[caminho.suffix] = resposta.json()["id"]

    if len(enviados) < 2:
        return

    # A ingestão é assíncrona; sem esperar, a busca acha um índice vazio e o
    # teste acusaria um erro que não existe.
    for _ in range(30):
        arquivos = client.get("/api/files", headers=auth(token)).json()
        prontos = [
            a
            for a in arquivos
            if a["id"] in enviados.values() and (a.get("status") or "") in ("ready", "pronto")
        ]
        if len(prontos) == 2:
            break
        time.sleep(6)
    check("documentos", "ingestão concluída nos dois arquivos", len(prontos) == 2, str(prontos))

    # Vincular ao agente e publicar: sem isso o RAG não enxerga o arquivo.
    versao = client.get(f"/api/templates/{template_id}", headers=auth(token)).json()
    ativa = versao.get("active_version") or {}
    agentes = []
    for a in ativa.get("agents", []):
        copia = dict(a)
        if copia["name"] == "gerais":
            copia["file_ids"] = list(enviados.values())
            copia["tools"] = sorted(set((copia.get("tools") or []) + ["query_agent_rag"]))
        agentes.append(
            {
                k: copia.get(k)
                for k in (
                    "name",
                    "description",
                    "prompt",
                    "tools",
                    "ai_service_id",
                    "model_override",
                    "reasoning_effort",
                    "datasource_ids",
                    "file_ids",
                )
            }
        )
    nova = client.post(
        f"/api/templates/{template_id}/versions",
        json={
            "supervisor_prompt": ativa.get("supervisor_prompt", ""),
            "supervisor_ai_service_id": ativa.get("supervisor_ai_service_id"),
            "max_steps": ativa.get("max_steps", 18),
            "history_limit": ativa.get("history_limit", 100),
            "compress_history": ativa.get("compress_history", False),
            "agents": agentes,
            "datasource_ids": ativa.get("datasource_ids", []),
        },
        headers=auth(token),
        timeout=300.0,
    )
    publicada = check(
        "documentos", "versão com documentos publicada", nova.status_code == 201, nova.text
    )
    if not publicada:
        return
    client.post(
        f"/api/templates/{template_id}/deploy",
        json={"version_id": nova.json()["id"]},
        headers=auth(token),
        timeout=300.0,
    )

    r = enviar(
        client,
        token,
        "Consultando nossos documentos internos: qual é o código interno do "
        "contrato guarda-chuva?",
        template_id,
    )
    check(
        "documentos",
        "agente lê o .txt e traz o código exato",
        "LKT-9931-ZX" in (r["reply"] or ""),
        r["reply"][:300],
    )

    r = enviar(
        client,
        token,
        "E qual é a meta de faturamento para 2027, segundo o documento de metas?",
        template_id,
    )
    check(
        "documentos",
        "agente lê o .docx e traz o valor exato",
        "48.750.000" in (r["reply"] or "") or "48750000" in (r["reply"] or ""),
        r["reply"][:300],
    )

    r = enviar(
        client,
        token,
        "Qual é o código interno do contrato de manutenção predial?",
        template_id,
    )
    inventou = "LKT" in (r["reply"] or "").upper() and "9931" in (r["reply"] or "")
    check(
        "documentos",
        "não inventa código para documento que não existe",
        not inventou,
        r["reply"][:300],
    )


def suite_artefatos(client, token, template_id):
    """Não basta o agente dizer que gerou: o arquivo tem que existir."""
    print("\n-- artefatos: o arquivo existe mesmo? --")

    r = enviar(
        client,
        token,
        "Traga a despesa em educação dos 5 maiores municípios de Rondônia no "
        "ano mais recente e faça um gráfico de barras.",
        template_id,
    )
    chat_id = r["chat_id"]
    tipos = [a.get("kind") or a.get("type") for a in r["artifacts"]]
    check("artefatos", "gráfico chegou como artefato", "chart" in tipos, str(tipos))

    r2 = enviar(client, token, "Exporte esses mesmos dados em planilha.", template_id, chat_id)
    r3 = enviar(client, token, "Agora gere um PDF com esse resumo.", template_id, chat_id)

    listagem = client.get(f"/api/chats/{chat_id}/artifacts", headers=auth(token)).json()
    listagem = listagem.get("items", listagem) if isinstance(listagem, dict) else listagem
    por_tipo = {}
    for a in listagem:
        por_tipo.setdefault(a.get("kind"), []).append(a)
    check(
        "artefatos",
        "chat tem gráfico, planilha e documento",
        "chart" in por_tipo and "file" in por_tipo,
        str(list(por_tipo)),
    )

    # O teste que importa: baixar. Artefato listado e vazio já aconteceu.
    baixados = 0
    for a in listagem[:8]:
        resposta = client.get(
            f"/api/artifacts/{a['id']}/download", headers=auth(token), timeout=300.0
        )
        if resposta.status_code == 200 and len(resposta.content) > 200:
            baixados += 1
    check(
        "artefatos",
        "artefatos baixam com conteúdo",
        baixados >= 2,
        f"{baixados} de {len(listagem[:8])} baixaram com corpo",
    )

    grafico = (por_tipo.get("chart") or [None])[0]
    if grafico:
        payload = client.get(
            f"/api/artifacts/{grafico['id']}/payload", headers=auth(token), timeout=300.0
        )
        corpo = payload.text
        check(
            "artefatos",
            "gráfico traz dados de verdade (Plotly com pontos)",
            payload.status_code == 200 and ('"data"' in corpo and len(corpo) > 300),
            corpo[:200],
        )
    check(
        "artefatos",
        "planilha e PDF sem erro no turno",
        not r2["error"] and not r3["error"],
    )
    return chat_id


def suite_memoria(client, token, template_id):
    print("\n-- memória: guarda e lembra em outra conversa --")
    marca = uuid.uuid4().hex[:6].upper()
    fato = f"O nosso código de campanha deste trimestre é CMP-{marca}."

    primeiro = enviar(
        client,
        token,
        f"Guarde esta informação para as próximas conversas: {fato}",
        template_id,
    )
    check("memória", "turno de gravação sem erro", not primeiro["error"])

    # A extração roda depois do turno, fora do caminho da resposta.
    encontrada = False
    for _ in range(20):
        time.sleep(6)
        memorias = client.get("/api/memories", headers=auth(token)).json()
        memorias = memorias.get("items", memorias) if isinstance(memorias, dict) else memorias
        if any(marca in (m.get("content") or "") for m in memorias):
            encontrada = True
            break
    check("memória", "fato virou memória", encontrada, f"procurando CMP-{marca}")

    # Conversa nova: sem chat_id, o histórico não ajuda — só a memória.
    segundo = enviar(client, token, "Qual é o código de campanha deste trimestre?", template_id)
    check(
        "memória",
        "lembra numa conversa nova",
        marca in (segundo["reply"] or ""),
        segundo["reply"][:300],
    )
    return marca


def suite_isolamento(client, token_a, marca, chat_com_artefato, artefato_id):
    """A segunda empresa não pode ver nada da primeira."""
    print("\n-- isolamento entre empresas --")
    master = client.post(
        "/api/auth/login", json={"email": MASTER_EMAIL, "password": MASTER_SENHA}
    ).json()["token"]
    chave = f"vizinha-{uuid.uuid4().hex[:6]}"
    tenant = client.post(
        "/api/tenants",
        json={"name": f"Vizinha {chave}", "tenant_key": chave},
        headers=auth(master),
    ).json()
    senha = "vizinha-senha-forte-123"
    client.post(
        "/api/users",
        json={
            "tenant_id": tenant["id"],
            "name": "Dona Vizinha",
            "email": f"dono@{chave}.com",
            "password": senha,
        },
        headers=auth(master),
    )
    token_b = client.post(
        "/api/auth/login", json={"email": f"dono@{chave}.com", "password": senha}
    ).json()["token"]

    memorias = client.get("/api/memories", headers=auth(token_b)).json()
    memorias = memorias.get("items", memorias) if isinstance(memorias, dict) else memorias
    check(
        "isolamento",
        "não vê memória da outra empresa",
        not any(marca in (m.get("content") or "") for m in memorias),
        f"{len(memorias)} memórias visíveis",
    )

    arquivos = client.get("/api/files", headers=auth(token_b)).json()
    arquivos = arquivos.get("items", arquivos) if isinstance(arquivos, dict) else arquivos
    check(
        "isolamento",
        "não vê documentos da outra empresa",
        len(arquivos) == 0,
        str(len(arquivos)),
    )

    resposta = client.get(f"/api/chats/{chat_com_artefato}/artifacts", headers=auth(token_b))
    corpo = resposta.text
    check(
        "isolamento",
        "não lista artefatos de chat alheio",
        resposta.status_code in (403, 404) or corpo.strip() in ("[]", '{"items":[]}'),
        f"{resposta.status_code} {corpo[:150]}",
    )

    if artefato_id:
        resposta = client.get(f"/api/artifacts/{artefato_id}/download", headers=auth(token_b))
        check(
            "isolamento",
            "não baixa artefato alheio",
            resposta.status_code in (403, 404),
            str(resposta.status_code),
        )

    templates = client.get("/api/templates", headers=auth(token_b)).json()
    check(
        "isolamento",
        "não vê templates da outra empresa",
        len(templates) == 0,
        str(len(templates)),
    )


def main() -> int:
    pasta = Path(os.environ.get("LICITA_FIXTURES", "docs/fixtures-qa"))
    pasta.mkdir(parents=True, exist_ok=True)

    with httpx.Client(base_url=BACKEND, timeout=600.0) as client:
        token = client.post(
            "/api/auth/login", json={"email": EMAIL, "password": SENHA}
        ).json()["token"]
        templates = client.get("/api/templates", headers=auth(token)).json()
        template_id = next(t for t in templates if t["name"] == TEMPLATE)["id"]

        suite_documentos(client, token, template_id, pasta)
        chat_id = suite_artefatos(client, token, template_id)
        marca = suite_memoria(client, token, template_id)

        artefatos = client.get(f"/api/chats/{chat_id}/artifacts", headers=auth(token)).json()
        artefatos = artefatos.get("items", artefatos) if isinstance(artefatos, dict) else artefatos
        artefato_id = artefatos[0]["id"] if artefatos else None
        suite_isolamento(client, token, marca, chat_id, artefato_id)

    destino = Path(__file__).parent.parent / "docs" / "qa-memoria-artefatos.json"
    destino.write_text(json.dumps(resultados, ensure_ascii=False, indent=2), encoding="utf-8")
    falhas = [r for r in resultados if not r["ok"]]
    print(f"\n{'=' * 70}")
    print(f"verificações: {len(resultados)} | falhas: {len(falhas)}")
    for f in falhas:
        print(f"  FALHA {f['grupo']}: {f['nome']} — {f['detalhe'][:160]}")
    print(f"relatório: {destino}")
    print("RESULTADO:", "APROVADO" if not falhas else "REPROVADO")
    return 1 if falhas else 0


if __name__ == "__main__":
    sys.exit(main())
