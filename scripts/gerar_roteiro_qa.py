"""Gera o roteiro de QA em .docx para o time de qualidade.

Por que um gerador e não um .docx escrito à mão: o roteiro precisa acompanhar
o sistema. Quando uma tela muda de nome ou um campo some, editar o script e
rodar de novo é mais barato — e mais confiável — do que caçar o parágrafo
dentro de um binário do Word.

Uso:
    python scripts/gerar_roteiro_qa.py [caminho-de-saida.docx]
"""

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

AZUL = RGBColor(0x1F, 0x3A, 0x5F)
CINZA = RGBColor(0x55, 0x5F, 0x6D)
VERMELHO = RGBColor(0xB4, 0x23, 0x18)
VERDE = RGBColor(0x1B, 0x6E, 0x3C)

BACKEND = "https://teste-ia-backend-x27vtpiida-uc.a.run.app"
CHATWOOT = "https://chatwoot-web-x27vtpiida-uc.a.run.app"


# --------------------------------------------------------------------------
# Blocos de montagem
# --------------------------------------------------------------------------


def _sombrear(celula, cor_hex: str) -> None:
    elemento = OxmlElement("w:shd")
    elemento.set(qn("w:val"), "clear")
    elemento.set(qn("w:fill"), cor_hex)
    celula._tc.get_or_add_tcPr().append(elemento)


def _borda_caixa(tabela) -> None:
    bordas = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        linha = OxmlElement(f"w:{lado}")
        linha.set(qn("w:val"), "single")
        linha.set(qn("w:sz"), "6")
        linha.set(qn("w:color"), "BFC7D2")
        bordas.append(linha)
    tabela._tbl.tblPr.append(bordas)


class Roteiro:
    def __init__(self) -> None:
        self.doc = Document()
        self._pagina()
        self._estilos()
        self.numero_caso = 0

    def _pagina(self) -> None:
        for secao in self.doc.sections:
            secao.top_margin = Cm(2)
            secao.bottom_margin = Cm(2)
            secao.left_margin = Cm(2.2)
            secao.right_margin = Cm(2.2)

    def _estilos(self) -> None:
        normal = self.doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15

    # -- texto ------------------------------------------------------------

    def titulo_capa(self, texto: str, subtitulo: str) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(texto)
        r.font.size = Pt(30)
        r.font.bold = True
        r.font.color.rgb = AZUL
        s = self.doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rs = s.add_run(subtitulo)
        rs.font.size = Pt(13)
        rs.font.color.rgb = CINZA

    def h1(self, texto: str, quebrar: bool = True) -> None:
        if quebrar:
            self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        p = self.doc.add_heading(texto, level=1)
        for r in p.runs:
            r.font.color.rgb = AZUL

    def h2(self, texto: str) -> None:
        p = self.doc.add_heading(texto, level=2)
        for r in p.runs:
            r.font.color.rgb = AZUL

    def h3(self, texto: str) -> None:
        p = self.doc.add_heading(texto, level=3)
        for r in p.runs:
            r.font.color.rgb = CINZA

    def p(self, texto: str, italico: bool = False, negrito: bool = False) -> None:
        par = self.doc.add_paragraph()
        r = par.add_run(texto)
        r.italic = italico
        r.bold = negrito

    def bullet(self, texto: str) -> None:
        self.doc.add_paragraph(texto, style="List Bullet")

    def passo(self, texto: str) -> None:
        self.doc.add_paragraph(texto, style="List Number")

    def codigo(self, texto: str) -> None:
        par = self.doc.add_paragraph()
        par.paragraph_format.left_indent = Cm(0.6)
        par.paragraph_format.space_after = Pt(8)
        r = par.add_run(texto)
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)

    def aviso(self, titulo: str, texto: str, cor: str = "FFF4E5") -> None:
        tabela = self.doc.add_table(rows=1, cols=1)
        tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
        celula = tabela.cell(0, 0)
        _sombrear(celula, cor)
        _borda_caixa(tabela)
        par = celula.paragraphs[0]
        r = par.add_run(f"{titulo}\n")
        r.bold = True
        par.add_run(texto)
        self.doc.add_paragraph()

    def tabela(self, cabecalho: list[str], linhas: list[list[str]], larguras=None) -> None:
        tabela = self.doc.add_table(rows=1, cols=len(cabecalho))
        _borda_caixa(tabela)
        for i, nome in enumerate(cabecalho):
            celula = tabela.rows[0].cells[i]
            celula.text = ""
            r = celula.paragraphs[0].add_run(nome)
            r.bold = True
            r.font.size = Pt(9.5)
            _sombrear(celula, "E8EDF4")
        for linha in linhas:
            celulas = tabela.add_row().cells
            for i, valor in enumerate(linha):
                celulas[i].text = ""
                r = celulas[i].paragraphs[0].add_run(str(valor))
                r.font.size = Pt(9.5)
        if larguras:
            for linha in tabela.rows:
                for i, largura in enumerate(larguras):
                    linha.cells[i].width = Cm(largura)
        self.doc.add_paragraph()

    # -- blocos de teste --------------------------------------------------

    def caso(self, codigo: str, titulo: str, objetivo: str) -> None:
        self.numero_caso += 1
        self.h3(f"{codigo} — {titulo}")
        par = self.doc.add_paragraph()
        r = par.add_run("Por que este teste existe: ")
        r.bold = True
        par.add_run(objetivo)

    def como_fazer(self, passos: list[str]) -> None:
        self.p("Como fazer:", negrito=True)
        for texto in passos:
            self.passo(texto)

    def validar(self, itens: list[str]) -> None:
        self.p("O que precisa acontecer (se não acontecer, é bug):", negrito=True)
        for item in itens:
            self.bullet(item)

    def evidencia(self, instrucao: str = "Cole aqui o print da tela.") -> None:
        tabela = self.doc.add_table(rows=2, cols=1)
        _borda_caixa(tabela)
        topo = tabela.cell(0, 0)
        _sombrear(topo, "F2F5F9")
        topo.text = ""
        r = topo.paragraphs[0].add_run(f"EVIDÊNCIA — {instrucao}")
        r.bold = True
        r.font.size = Pt(9)
        corpo = tabela.cell(1, 0)
        corpo.text = ""
        for _ in range(6):
            corpo.add_paragraph()
        self.doc.add_paragraph()

    def resultado(self) -> None:
        self.tabela(
            ["Resultado", "Data / hora", "Testado por", "Observações (o que viu de estranho)"],
            [["(  ) Passou   (  ) Falhou   (  ) Bloqueado", "", "", ""]],
            larguras=[5.0, 2.8, 2.8, 6.4],
        )

    def bloco(
        self,
        codigo: str,
        titulo: str,
        objetivo: str,
        passos: list[str],
        validacoes: list[str],
        evidencia: str = "Cole aqui o print da tela.",
    ) -> None:
        self.caso(codigo, titulo, objetivo)
        self.como_fazer(passos)
        self.validar(validacoes)
        self.evidencia(evidencia)
        self.resultado()

    def salvar(self, caminho: Path) -> None:
        self.doc.save(caminho)


# --------------------------------------------------------------------------
# Conteúdo
# --------------------------------------------------------------------------


def montar() -> Roteiro:
    r = Roteiro()

    # ---------------------------------------------------------------- capa
    for _ in range(6):
        r.doc.add_paragraph()
    r.titulo_capa(
        "Roteiro de QA — Plataforma de Agentes de IA",
        "Guia completo de testes, do zero ao atendimento em produção",
    )
    r.doc.add_paragraph()
    r.tabela(
        ["Campo", "Preencher"],
        [
            ["Responsável pelos testes", ""],
            ["Período de execução", ""],
            ["Versão / data do sistema testado", ""],
            ["Ambiente usado", ""],
            ["Revisado por", ""],
        ],
        larguras=[6.0, 11.0],
    )
    r.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ------------------------------------------------- como usar / conceitos
    r.h1("1. Como usar este documento", quebrar=False)
    r.p(
        "Este roteiro foi escrito para quem está chegando agora. Ele não assume "
        "que você já conhece o sistema: cada seção começa explicando o que a "
        "funcionalidade faz e por que ela existe, e só depois pede que você "
        "teste. Leia a explicação antes de executar — testar sem entender "
        "encontra menos problemas, e você vai precisar desse entendimento "
        "depois para treinar e demonstrar o produto."
    )
    r.p("Cada teste segue sempre o mesmo formato:", negrito=True)
    r.tabela(
        ["Parte", "O que é"],
        [
            [
                "Por que este teste existe",
                "O risco real que ele cobre. É o motivo de você estar fazendo isso.",
            ],
            ["Como fazer", "Passo a passo. Siga na ordem; a ordem importa."],
            [
                "O que precisa acontecer",
                "O resultado esperado. Se der diferente, é bug — registre.",
            ],
            ["Evidência", "Espaço para colar o print. Print é a prova de que o teste rodou."],
            ["Resultado", "Passou / Falhou / Bloqueado, mais o que você observou."],
        ],
        larguras=[5.0, 12.0],
    )
    r.p(
        "Use “Bloqueado” quando você não conseguiu nem executar o teste — por "
        "exemplo, a tela não abriu, ou um teste anterior era pré-requisito e "
        "falhou. Bloqueado não é a mesma coisa que Falhou, e essa diferença "
        "ajuda muito a entender o que aconteceu."
    )

    r.aviso(
        "Regra de ouro do print",
        "Tire o print da tela inteira, não só do pedaço que interessa. Muita "
        "informação útil para diagnosticar um problema está na barra de "
        "endereço, no menu lateral e nas mensagens de canto de tela. Um print "
        "recortado costuma esconder exatamente o que o desenvolvedor precisava ver.",
    )

    r.aviso(
        "Quando algo falhar, não pare",
        "Registre, tire o print e siga para o próximo teste, a menos que o "
        "teste seguinte dependa daquele que quebrou. Um roteiro interrompido no "
        "meio entrega muito menos informação do que um roteiro completo com "
        "cinco itens marcados como falha.",
        cor="FDECEA",
    )

    # ------------------------------------------------------------ glossário
    r.h1("2. Glossário — o vocabulário do sistema")
    r.p(
        "Estes termos aparecem o tempo todo, tanto no produto quanto nas "
        "conversas com o time. Vale ler com calma uma vez; depois volte aqui "
        "sempre que precisar."
    )
    r.tabela(
        ["Termo", "O que significa na prática"],
        [
            [
                "Empresa (tenant)",
                "Um cliente da plataforma. Cada empresa é uma caixa fechada: seus "
                "usuários, dados, chaves e conversas não podem ser vistos por outra empresa. "
                "Boa parte deste roteiro existe para provar que essa caixa não vaza.",
            ],
            [
                "Master",
                "O administrador da plataforma inteira (você/o dono). Vê e cria "
                "empresas. Não é usuário de nenhuma empresa específica.",
            ],
            [
                "Perfil",
                "Um conjunto de permissões (ver, criar, editar, apagar) por recurso. "
                "É o perfil que decide quais menus a pessoa enxerga.",
            ],
            [
                "Template",
                "A configuração de um agente de IA: o que ele sabe, com que tom "
                "responde, que ferramentas pode usar e a quais bancos tem acesso. "
                "Uma empresa pode ter vários — por exemplo um para vendas e outro para suporte.",
            ],
            [
                "Versão do template",
                "Cada alteração publicada vira uma versão nova e imutável. A versão "
                "antiga não muda. Isso permite voltar atrás sem medo.",
            ],
            [
                "Supervisor e especialistas",
                "Dentro de um template, o supervisor é quem conversa com o cliente. "
                "Ele aciona os especialistas como se fossem ferramentas. O cliente só "
                "vê a resposta final.",
            ],
            [
                "Ferramenta (tool)",
                "Uma ação que o agente pode executar: consultar banco, gerar gráfico, "
                "emitir cobrança, buscar na web. O agente só usa o que o template autorizou.",
            ],
            [
                "Fonte de dados (datasource)",
                "A conexão com um banco de dados do cliente. A leitura é livre; a "
                "escrita é restrita a tabelas autorizadas e pode exigir confirmação.",
            ],
            [
                "Artefato",
                "Um arquivo produzido durante a conversa: gráfico, planilha, PDF.",
            ],
            [
                "Memória",
                "Fatos que o agente guarda entre conversas diferentes (por exemplo, "
                "a preferência de um cliente). Não confunda com o histórico da conversa atual.",
            ],
            [
                "Serviço de IA",
                "De onde vem o modelo que responde. Pode ser uma conta conectada da "
                "empresa, ou um combo que reveza entre várias contas.",
            ],
            [
                "Combo",
                "Um grupo de modelos que se revezam. Serve para não estourar o limite "
                "de uso de uma conta só.",
            ],
            [
                "W-API",
                "O serviço que conecta um número de WhatsApp à plataforma.",
            ],
            [
                "Chatwoot",
                "O painel onde atendentes humanos respondem conversas. A IA atende "
                "primeiro e passa para o humano quando precisa.",
            ],
            [
                "Inbox (caixa de entrada)",
                "Dentro do Chatwoot, um canal de atendimento — por exemplo, um número de WhatsApp.",
            ],
            [
                "Handoff",
                "O momento em que a IA sai de cena e entrega a conversa para um atendente humano.",
            ],
            [
                "PIX / Mercado Pago",
                "A cobrança que o agente consegue emitir dentro da conversa.",
            ],
        ],
        larguras=[4.2, 12.8],
    )

    # ---------------------------------------------------------- ambiente
    r.h1("3. Antes de começar — ambiente e acessos")
    r.h2("3.1 Endereços")
    r.tabela(
        ["O quê", "Endereço"],
        [
            ["Plataforma (onde você vai passar 80% do tempo)", BACKEND],
            ["Painel de atendimento (Chatwoot)", CHATWOOT],
            [
                "Instância de IA de uma empresa (exemplo)",
                "https://ia-<chave-da-empresa>.rangeltech.net",
            ],
        ],
        larguras=[7.0, 10.0],
    )

    r.h2("3.2 Credenciais — peça ao responsável e anote aqui")
    r.aviso(
        "Não compartilhe este documento preenchido",
        "Depois de anotar as senhas abaixo, este arquivo passa a ser um documento "
        "sensível. Não mande por grupo de WhatsApp, não suba em drive público e "
        "não deixe aberto em máquina compartilhada. Se precisar enviar o roteiro "
        "para alguém, envie a versão em branco.",
        cor="FDECEA",
    )
    r.tabela(
        ["Acesso", "Usuário", "Senha", "Para que serve"],
        [
            ["Master da plataforma", "", "", "Criar e apagar empresas"],
            ["Usuário administrador da Empresa A", "", "", "Testes do dia a dia"],
            ["Usuário administrador da Empresa B", "", "", "Testes de isolamento"],
            ["Usuário administrador da Empresa C", "", "", "Testes simultâneos"],
            ["Usuário sem permissões (limitado)", "", "", "Testar bloqueio de menu"],
            ["Painel W-API", "", "", "Conectar números de WhatsApp"],
            ["Mercado Pago (sandbox)", "", "", "Testar cobrança PIX"],
        ],
        larguras=[4.6, 3.6, 3.6, 5.2],
    )

    r.h2("3.3 O que você precisa ter em mãos")
    r.bullet("Um computador com Chrome ou Edge atualizado.")
    r.bullet("Um celular com WhatsApp para escanear QR Code e mandar mensagens de teste.")
    r.bullet("Um segundo número de WhatsApp (pode ser de um colega) para simular o cliente.")
    r.bullet("Acesso ao painel da W-API com pelo menos duas instâncias disponíveis.")
    r.bullet("Uma conta Mercado Pago em modo sandbox (teste), com as credenciais de teste.")
    r.bullet("Bancos de dados de teste — a seção 9 explica como conseguir cada um.")

    r.h2("3.4 Ordem recomendada")
    r.p(
        "Execute as seções na ordem numérica. Elas foram encadeadas de "
        "propósito: a seção 5 cria as empresas que a seção 12 vai usar para "
        "testar atendimento simultâneo. Pular seções costuma gerar “bugs” que "
        "na verdade são pré-requisitos faltando."
    )

    # =====================================================================
    r.h1("4. Suíte 1 — Acesso, login e permissões")
    r.p(
        "Esta é a porta de entrada do sistema. Um problema aqui afeta todo o "
        "resto, então ela vem primeiro. O ponto mais importante desta suíte é "
        "entender que o menu que aparece na tela é montado a partir das "
        "permissões do perfil do usuário — não é igual para todo mundo."
    )

    r.bloco(
        "1.1",
        "Login com credenciais corretas",
        "Se o login falhar, nada mais pode ser testado.",
        [
            f"Abra {BACKEND} no navegador.",
            "Digite o e-mail e a senha do administrador da Empresa A.",
            "Clique em Entrar.",
        ],
        [
            "A tela carrega sem erro e mostra o nome da empresa no topo.",
            "O menu superior aparece com as opções permitidas ao perfil.",
            "Não aparece nenhuma mensagem vermelha de erro.",
        ],
        "Cole o print da tela inicial logo após o login.",
    )

    r.bloco(
        "1.2",
        "Login com senha errada",
        "O sistema precisa recusar o acesso e dizer isso com clareza, sem "
        "revelar se o e-mail existe (isso ajudaria quem tenta invadir).",
        [
            "Saia do sistema (botão Sair, no canto superior direito).",
            "Tente entrar com o e-mail correto e uma senha inventada.",
        ],
        [
            "O acesso é recusado.",
            "Aparece uma mensagem de erro compreensível, em português.",
            "A mensagem NÃO diz algo como “senha incorreta para este e-mail” — "
            "ela não deve confirmar que o e-mail existe.",
            "Você continua na tela de login.",
        ],
    )

    r.bloco(
        "1.3",
        "Login como master",
        "O master enxerga a plataforma inteira e tem um menu diferente. "
        "Confundir os dois papéis é a origem de muita confusão no suporte.",
        [
            "Saia e entre com as credenciais do master.",
            "Observe o menu superior com atenção.",
        ],
        [
            "Aparece o menu Empresas.",
            "NÃO aparecem os menus Início, Chat e Memórias — o master não "
            "conversa com agentes, ele administra empresas.",
            "A diferença em relação ao print do teste 1.1 é visível.",
        ],
        "Cole o print do menu do master, para comparar com o do teste 1.1.",
    )

    r.bloco(
        "1.4",
        "Usuário com permissões limitadas vê menos menus",
        "O menu é gerado pelas permissões. Se um usuário limitado enxergar um "
        "menu que não deveria, isso é uma falha de segurança, não um detalhe visual.",
        [
            "Entre como administrador da Empresa A.",
            "Vá em Perfis e crie um perfil chamado “Somente Chat”, marcando apenas "
            "a permissão de visualizar chat.",
            "Vá em Usuários e crie um usuário novo com esse perfil.",
            "Saia e entre com esse novo usuário.",
        ],
        [
            "O usuário limitado vê Início e Chat.",
            "Ele NÃO vê Usuários, Perfis, Templates, Fontes de dados, Pagamentos nem Integrações.",
            "Tentar abrir uma dessas telas digitando o endereço direto na barra "
            "(por exemplo, /usuarios) não mostra os dados — o sistema barra.",
        ],
        "Cole dois prints: o menu do usuário limitado e a tela ao tentar acessar /usuarios direto.",
    )

    r.bloco(
        "1.5",
        "Sessão encerrada ao sair",
        "Sair precisa realmente encerrar o acesso — não apenas esconder a tela.",
        [
            "Estando logado, clique em Sair.",
            "Aperte a seta de voltar do navegador.",
        ],
        [
            "Você não volta para dentro do sistema.",
            "O sistema pede login de novo.",
        ],
    )

    # =====================================================================
    r.h1("5. Suíte 2 — Empresas: criar, apagar e criar de novo")
    r.p(
        "Esta suíte é feita como o dono pediu: criar, apagar e criar de novo, "
        "para descobrir problemas que só aparecem no segundo cadastro — nome "
        "repetido, resíduo do cadastro anterior, chave que não foi liberada. "
        "É um dos testes mais produtivos que existem, porque quase todo sistema "
        "trata bem o primeiro cadastro e mal o segundo."
    )
    r.aviso(
        "Só o master faz isso",
        "Todos os testes desta suíte exigem login como master. Se o menu "
        "Empresas não aparecer, você está logado como usuário de empresa.",
    )

    r.bloco(
        "2.1",
        "Criar a Empresa A",
        "É o cadastro base de tudo que vem depois.",
        [
            "Entre como master e clique em Empresas.",
            "Clique no botão de criar empresa.",
            "Preencha o nome como “QA Empresa A” e a chave como “qa-empresa-a”.",
            "Salve.",
        ],
        [
            "A empresa aparece na lista imediatamente.",
            "A chave fica exatamente como digitada, em minúsculas e sem espaços.",
            "Nenhum erro aparece.",
        ],
        "Cole o print da lista de empresas com a Empresa A criada.",
    )

    r.bloco(
        "2.2",
        "Criar empresa com chave repetida",
        "Duas empresas com a mesma chave criariam ambiguidade em endereços e "
        "integrações. O sistema tem que recusar.",
        [
            "Tente criar outra empresa usando a mesma chave “qa-empresa-a”.",
        ],
        [
            "O sistema recusa.",
            "A mensagem explica o motivo em português claro.",
            "Nenhuma empresa duplicada aparece na lista.",
        ],
    )

    r.bloco(
        "2.3",
        "Criar o primeiro usuário da Empresa A",
        "O primeiro usuário de uma empresa recebe automaticamente o perfil de "
        "Administrador. Isso é intencional: sem isso, a empresa nasceria sem "
        "ninguém capaz de configurá-la — e o sintoma seria “sumiram os menus”.",
        [
            "Ainda como master, entre na Empresa A.",
            "Vá em Usuários e crie o usuário administrador (anote e-mail e senha na seção 3.2).",
            "Saia e entre com esse usuário.",
        ],
        [
            "O usuário entra normalmente.",
            "O perfil dele é Administrador.",
            "Ele enxerga o menu completo da empresa.",
        ],
        "Cole o print do menu completo desse primeiro usuário.",
    )

    r.bloco(
        "2.4",
        "Criar Empresa B e Empresa C",
        "Você vai precisar de três empresas para os testes de isolamento e de "
        "atendimento simultâneo mais adiante.",
        [
            "Repita os testes 2.1 e 2.3 para “QA Empresa B” (chave qa-empresa-b).",
            "Repita novamente para “QA Empresa C” (chave qa-empresa-c).",
        ],
        [
            "As três empresas aparecem na lista.",
            "Cada uma tem seu próprio usuário administrador.",
            "Entrar em uma não mostra dados da outra.",
        ],
        "Cole o print da lista com as três empresas.",
    )

    r.bloco(
        "2.5",
        "Apagar uma empresa",
        "Apagar é a operação mais perigosa do sistema. Precisa avisar antes e "
        "não pode deixar rastro visível depois.",
        [
            "Como master, crie uma empresa descartável chamada “QA Descartável”.",
            "Crie um usuário dentro dela.",
            "Agora apague essa empresa.",
        ],
        [
            "O sistema pede confirmação antes de apagar — não apaga no primeiro clique.",
            "Depois de confirmar, a empresa some da lista.",
            "O usuário que pertencia a ela não consegue mais entrar.",
        ],
        "Cole o print da confirmação e o print da lista já sem a empresa.",
    )

    r.bloco(
        "2.6",
        "Criar de novo com a mesma chave da empresa apagada",
        "Este é o teste que o dono pediu especificamente, e é onde sistemas "
        "costumam falhar: a chave da empresa apagada pode continuar “presa”, "
        "ou a empresa nova pode nascer enxergando dados da antiga.",
        [
            "Crie uma empresa nova usando exatamente a mesma chave da empresa que você apagou.",
            "Crie um usuário nela e entre.",
        ],
        [
            "A criação funciona OU o sistema explica claramente que a chave está reservada. "
            "Qualquer um dos dois é aceitável; erro genérico e sem explicação não é.",
            "Se foi criada: a empresa nova está completamente vazia — sem usuários "
            "antigos, sem templates, sem conversas, sem fontes de dados.",
            "Nada da empresa apagada reaparece.",
        ],
        "Cole o print do resultado — funcionando ou recusando com a mensagem.",
    )

    # =====================================================================
    r.h1("6. Suíte 3 — Serviços de IA: contas e combos")
    r.p(
        "Aqui é onde a empresa conecta as contas de IA que ela já paga. Existem "
        "três formas de conectar, e elas são diferentes de propósito, porque os "
        "provedores funcionam de formas diferentes:"
    )
    r.tabela(
        ["Forma", "Como funciona", "Exemplos"],
        [
            [
                "Chave de API",
                "Você cola uma chave copiada do site do provedor.",
                "Gemini, OpenAI, Anthropic, DeepSeek, Groq",
            ],
            [
                "Assinatura por redirecionamento",
                "Abre o site do provedor, você autoriza e volta com um código para colar.",
                "Claude, ChatGPT/Codex, Antigravity, Gemini CLI",
            ],
            [
                "Assinatura por código",
                "A tela mostra um código curto; você digita esse código no site do "
                "provedor e a tela percebe sozinha quando você confirma.",
                "GitHub Copilot, Qwen, Kimi, Kilo Code",
            ],
        ],
        larguras=[4.0, 8.0, 5.0],
    )
    r.aviso(
        "Pré-requisito importante",
        "Uma empresa só consegue conectar contas se tiver uma instância de IA "
        "provisionada. Isso é feito pelo responsável técnico, por script. Se a "
        "tela disser “Instância de IA ainda não provisionada”, isso NÃO é bug — "
        "peça o provisionamento antes de continuar esta suíte.",
    )

    r.bloco(
        "3.1",
        "Empresa sem instância mostra o aviso certo",
        "O cliente precisa entender o que fazer, em vez de ver uma tela quebrada.",
        [
            "Entre em uma empresa recém-criada (por exemplo a Empresa C).",
            "Vá em Serviços de IA.",
        ],
        [
            "Aparece o aviso “Instância de IA ainda não provisionada”.",
            "O texto orienta a pedir o provisionamento ao administrador.",
            "A tela não mostra erro técnico, código de erro nem tela em branco.",
        ],
    )

    r.bloco(
        "3.2",
        "Conectar conta por chave de API",
        "É o caminho mais simples e o mais usado. Também é onde se verifica que "
        "a chave não volta para a tela depois de salva.",
        [
            "Em uma empresa com instância provisionada, vá em Serviços de IA.",
            "Clique no card “Google Gemini”.",
            "Preencha o apelido (por exemplo, “Gemini da empresa”) e cole a chave.",
            "Clique em Conectar.",
        ],
        [
            "A conta aparece na lista abaixo com o apelido escolhido.",
            "O status aparece como Ativa ou Sem uso ainda.",
            "Ao reabrir a tela, a chave NÃO aparece — nem parcialmente, nem "
            "mascarada com o valor real.",
            "O card do Gemini passa a mostrar um contador com o número de contas.",
        ],
        "Cole o print da lista de contas com a conta recém-conectada.",
    )

    r.bloco(
        "3.3",
        "Conectar assinatura Claude (fluxo de redirecionamento)",
        "É o caso que motivou toda essa funcionalidade: usar a assinatura que a "
        "empresa já paga, que não fornece chave de API.",
        [
            "Clique no card “Claude (assinatura)”.",
            "Clique em Abrir autorização — uma aba nova abre no site do provedor.",
            "Faça login e autorize.",
            "Copie a URL inteira da barra de endereço da página de retorno.",
            "Volte à plataforma e cole essa URL no campo indicado.",
            "Clique em Concluir.",
        ],
        [
            "A aba de autorização abre sozinha.",
            "Colar a URL inteira funciona — você não precisa recortar o código à mão.",
            "A conta aparece na lista como assinatura.",
            "Se você colar algo inválido, aparece uma mensagem clara em vez de a tela travar.",
        ],
        "Cole o print do modal aberto e o print da conta já conectada.",
    )

    r.bloco(
        "3.4",
        "Conectar GitHub Copilot (fluxo de código)",
        "É o terceiro fluxo, e o que mais confunde quem nunca viu. Vale testar "
        "mesmo que a empresa não vá usar Copilot, porque valida o mecanismo.",
        [
            "Clique no card “GitHub Copilot”.",
            "Anote o código curto que aparece (formato parecido com ABCD-1234).",
            "Abra o endereço indicado e digite o código lá.",
            "Volte para a plataforma e aguarde sem fechar o modal.",
        ],
        [
            "O código aparece na tela, legível e grande.",
            "O endereço para digitar o código está visível.",
            "Depois que você confirma no site, a janela fecha sozinha em até um minuto — "
            "você não precisa apertar nada.",
            "A conta aparece na lista.",
        ],
        "Cole o print do modal com o código visível.",
    )

    r.bloco(
        "3.5",
        "Criar um combo",
        "O combo é o que faz o revezamento entre modelos. Sem ele, uma conta "
        "que bate no limite derruba o atendimento.",
        [
            "Na seção Combos, escreva o nome “Produção”.",
            "Marque pelo menos dois modelos da lista.",
            "Clique em Criar combo.",
        ],
        [
            "O combo aparece na lista, mostrando os modelos escolhidos.",
            "Aparece a marcação “serviço de IA publicado”.",
            "Indo em Templates, o combo aparece como opção de serviço de IA — "
            "com o nome “Combo: Produção”.",
        ],
        "Cole o print do combo criado e o print da lista de serviços de IA no editor de template.",
    )

    r.bloco(
        "3.6",
        "A lista de modelos reflete só as contas conectadas",
        "O catálogo do roteador tem centenas de modelos, mas só alguns funcionam "
        "de verdade — os que têm conta conectada. Se aparecer modelo demais, o "
        "cliente monta um combo que quebra no meio do atendimento.",
        [
            "Conte quantos modelos aparecem na lista de combos.",
            "Compare com as contas conectadas na seção acima.",
        ],
        [
            "Só aparecem modelos dos provedores que a empresa realmente conectou.",
            "Se a empresa só tem conta Gemini, só aparecem modelos Gemini.",
            "Não aparecem centenas de modelos de provedores desconhecidos.",
        ],
    )

    r.bloco(
        "3.7",
        "Revezamento entre contas do mesmo provedor",
        "É o recurso principal: distribuir o uso entre várias contas para não "
        "estourar o limite de nenhuma.",
        [
            "Conecte uma segunda conta do mesmo provedor (por exemplo, um segundo Gemini).",
            "Observe se aparece a seção “Revezamento entre contas”.",
            "Troque a opção para “Alternar a cada chamada”.",
        ],
        [
            "A seção de revezamento só aparece quando existem duas ou mais contas "
            "do mesmo provedor.",
            "A escolha é salva — ao recarregar a página, ela continua igual.",
        ],
        "Cole o print da seção de revezamento.",
    )

    r.bloco(
        "3.8",
        "Remover uma conta",
        "Remover precisa tirar a conta dos dois lados: da plataforma e da instância.",
        [
            "Remova uma das contas conectadas.",
            "Recarregue a página.",
        ],
        [
            "A conta some da lista e não volta ao recarregar.",
            "Os modelos daquele provedor somem da lista de combos, se não sobrou outra conta dele.",
            "Combos existentes que usavam aquele provedor continuam listados, "
            "mas anote se algum deles parar de responder — isso é informação útil.",
        ],
    )

    # =====================================================================
    r.h1("7. Suíte 4 — Templates: o cérebro do agente")
    r.p(
        "O template define como o agente se comporta. Uma empresa pode ter "
        "vários, e cada conversa escolhe qual usar. Esta suíte testa criação, "
        "versionamento e o fato de que versões publicadas não mudam sozinhas."
    )

    r.bloco(
        "4.1",
        "Criar o primeiro template",
        "É o mínimo para conseguir conversar com um agente.",
        [
            "Entre na Empresa A e vá em Templates.",
            "Crie um template chamado “Atendimento Vendas”.",
            "Escreva no prompt do supervisor: “Você atende clientes de uma loja de "
            "ferragens. Seja objetivo e cordial.”",
            "Crie um especialista chamado “consultor”, com as ferramentas de consultar "
            "banco marcadas.",
            "Salve e publique a versão.",
        ],
        [
            "O template aparece na lista.",
            "Ele mostra que tem uma versão ativa.",
            "Ele aparece como opção na tela de Chat.",
        ],
        "Cole o print do template criado e publicado.",
    )

    r.bloco(
        "4.2",
        "Criar um segundo template na MESMA empresa",
        "Uma empresa costuma ter mais de um agente — vendas e suporte, por "
        "exemplo. Eles não podem se misturar.",
        [
            "Ainda na Empresa A, crie “Suporte Técnico”.",
            "Use um prompt bem diferente, por exemplo: “Você é do suporte técnico. "
            "Responda sempre começando com a palavra SUPORTE.”",
            "Publique.",
            "Vá ao Chat e converse com cada um dos dois, na mesma sessão de trabalho.",
        ],
        [
            "Os dois templates aparecem na lista de escolha do Chat.",
            "O agente de suporte responde começando com SUPORTE; o de vendas, não.",
            "Trocar de template não mistura o comportamento nem o histórico.",
        ],
        "Cole prints das duas conversas lado a lado, mostrando a diferença de comportamento.",
    )

    r.bloco(
        "4.3",
        "Criar um terceiro template com ferramentas diferentes",
        "Serve para confirmar que o agente só usa o que foi autorizado.",
        [
            "Crie “Analista de Dados” com as ferramentas de gráfico e planilha marcadas, "
            "além da consulta ao banco.",
            "Publique e converse pedindo um gráfico.",
        ],
        [
            "O agente gera o gráfico.",
            "Peça um gráfico ao template “Atendimento Vendas” (que não tem essa ferramenta): "
            "ele NÃO deve gerar — deve explicar que não consegue, e não inventar um link falso.",
        ],
        "Cole o print do gráfico gerado e o print da recusa do outro template.",
    )

    r.bloco(
        "4.4",
        "Editar um template já publicado gera versão nova",
        "As versões são imutáveis. Isso é o que permite voltar atrás com segurança.",
        [
            "Abra “Atendimento Vendas” e mude alguma coisa no prompt.",
            "Publique novamente.",
            "Procure o histórico de versões.",
        ],
        [
            "Aparece uma versão nova; a anterior continua listada.",
            "A conversa nova usa o comportamento novo.",
            "Nada indica que a versão antiga foi alterada.",
        ],
    )

    r.bloco(
        "4.5",
        "Templates não vazam entre empresas",
        "Este é um teste de segurança, não de conveniência.",
        [
            "Entre na Empresa B.",
            "Vá em Templates.",
        ],
        [
            "Os templates da Empresa A NÃO aparecem.",
            "A lista está vazia ou só com os templates da própria Empresa B.",
        ],
        "Cole o print da tela de templates da Empresa B.",
    )

    # =====================================================================
    r.h1("8. Suíte 5 — Conversa com o agente (Chat)")
    r.p(
        "Aqui você testa o produto do ponto de vista de quem usa. Preste atenção "
        "não só na resposta certa, mas em como ela chega: se o texto aparece "
        "progressivamente, se dá para saber que o agente está trabalhando, e se "
        "um erro aparece de forma compreensível."
    )

    r.bloco(
        "5.1",
        "Primeira conversa",
        "Valida o caminho completo: sua mensagem sai, o agente pensa e a resposta volta.",
        [
            "Vá em Chat, escolha o template “Atendimento Vendas”.",
            "Escreva “Olá, o que você consegue fazer?” e envie.",
        ],
        [
            "A resposta aparece aos poucos, palavra por palavra, e não de uma vez "
            "só depois de muito tempo.",
            "Enquanto o agente trabalha, existe alguma indicação visual disso.",
            "A resposta faz sentido com o prompt configurado.",
        ],
        "Cole o print da conversa.",
    )

    r.bloco(
        "5.2",
        "Conversa longa e histórico",
        "O agente precisa lembrar o que foi dito antes dentro da mesma conversa.",
        [
            "Diga “Meu nome é Carlos e eu prefiro ser chamado de Sr. Carlos”.",
            "Faça mais três perguntas quaisquer.",
            "Depois pergunte: “Como você deve me chamar?”",
        ],
        [
            "O agente responde “Sr. Carlos”.",
            "O histórico completo continua visível ao rolar para cima.",
        ],
        "Cole o print mostrando a pergunta inicial e a resposta final.",
    )

    r.bloco(
        "5.3",
        "Layout em tela cheia e em celular",
        "O chat é a tela mais usada. Ela precisa funcionar maximizada e no celular.",
        [
            "Maximize a janela do navegador em um monitor grande.",
            "Depois aperte F12, ative o modo dispositivo e teste em tamanho de celular.",
            "Teste também com a janela em metade da tela.",
        ],
        [
            "Em tela cheia, a área da conversa ocupa o espaço disponível — não fica "
            "espremida numa coluna estreita no meio.",
            "No celular, o menu vira um botão e a caixa de escrever fica acessível.",
            "Não aparece barra de rolagem horizontal em nenhum tamanho.",
        ],
        "Cole três prints: tela cheia, meia tela e celular.",
    )

    r.bloco(
        "5.4",
        "Anexar um arquivo à conversa",
        "Clientes mandam PDF e imagem o tempo todo.",
        [
            "Anexe um PDF de poucas páginas à conversa.",
            "Peça um resumo do conteúdo.",
        ],
        [
            "O upload mostra progresso ou confirmação.",
            "O agente responde sobre o conteúdo real do arquivo, não sobre o nome dele.",
        ],
        "Cole o print da conversa com o anexo e a resposta.",
    )

    r.bloco(
        "5.5",
        "Comportamento quando algo dá errado",
        "Erro vai acontecer. O que não pode é o cliente ficar sem saber o que houve.",
        [
            "Faça uma pergunta que exija banco de dados em um template que não tem "
            "nenhuma fonte configurada.",
        ],
        [
            "O agente responde explicando que não tem acesso àquela informação.",
            "NÃO aparece mensagem técnica crua na tela (nome de tabela, stack "
            "trace, código de erro).",
            "A conversa continua funcionando depois — não trava.",
        ],
    )

    # =====================================================================
    r.h1("9. Suíte 6 — Fontes de dados: um exemplo de cada tecnologia")
    r.p(
        "O sistema conecta quatro tipos de banco. Testar um de cada é importante "
        "porque cada um tem um jeito diferente de conectar e de reagir a erro. "
        "A tabela abaixo mostra o que cada um pede."
    )
    r.tabela(
        ["Tecnologia", "O que o formulário pede", "Onde conseguir para o teste"],
        [
            [
                "PostgreSQL",
                "Host, porta (5432), banco, usuário, senha",
                "Peça ao responsável um banco de teste, ou use o de demonstração",
            ],
            [
                "MySQL / MariaDB",
                "Host, porta (3306), banco, usuário, senha",
                "Serviço gratuito de teste na nuvem, ou instalação local",
            ],
            [
                "Google BigQuery",
                "Projeto e dataset",
                "Peça o projeto de teste ao responsável técnico",
            ],
            [
                "SQLite",
                "Caminho do arquivo",
                "Um arquivo .db pequeno, criado só para o teste",
            ],
        ],
        larguras=[3.4, 6.2, 7.4],
    )

    r.bloco(
        "6.1",
        "Conectar PostgreSQL e testar",
        "É a tecnologia mais usada pelos clientes.",
        [
            "Vá em Fontes de dados e clique para criar.",
            "Escolha PostgreSQL e preencha os dados.",
            "Salve e clique no botão de testar conexão.",
        ],
        [
            "O teste retorna sucesso.",
            "Ao reabrir o cadastro, a senha NÃO aparece.",
            "A fonte aparece na lista.",
        ],
        "Cole o print do teste bem-sucedido.",
    )

    r.bloco(
        "6.2",
        "Conectar MySQL / MariaDB",
        "Confirma que a porta padrão muda sozinha e que a conexão funciona.",
        [
            "Crie uma fonte escolhendo MySQL.",
            "Repare se a porta sugerida muda para 3306.",
            "Preencha e teste.",
        ],
        [
            "A porta sugerida vira 3306 ao escolher MySQL.",
            "O teste retorna sucesso.",
        ],
    )

    r.bloco(
        "6.3",
        "Conectar Google BigQuery",
        "BigQuery não usa host e senha, e sim projeto e dataset. O formulário "
        "precisa mudar de acordo.",
        [
            "Crie uma fonte escolhendo BigQuery.",
            "Repare que os campos mudam: aparecem Projeto e Dataset.",
            "Preencha e teste.",
        ],
        [
            "Os campos de host, porta e senha somem ao escolher BigQuery.",
            "O teste retorna sucesso.",
        ],
        "Cole o print do formulário do BigQuery, mostrando os campos diferentes.",
    )

    r.bloco(
        "6.4",
        "Conectar SQLite",
        "É o caso mais simples: um arquivo.",
        [
            "Crie uma fonte escolhendo SQLite e informe o caminho do arquivo.",
            "Teste.",
        ],
        ["O teste retorna sucesso.", "Aparece só o campo de caminho do arquivo."],
    )

    r.bloco(
        "6.5",
        "Conexão com dados errados",
        "Errar a senha é o erro mais comum do mundo real. A mensagem precisa ajudar.",
        [
            "Crie uma fonte PostgreSQL com a senha errada de propósito.",
            "Teste.",
        ],
        [
            "O teste falha, e isso fica claro na tela.",
            "A mensagem indica que houve falha de conexão.",
            "A mensagem NÃO expõe a senha nem o texto interno do erro do banco.",
        ],
        "Cole o print da mensagem de erro.",
    )

    r.bloco(
        "6.6",
        "O agente consulta o banco de verdade",
        "De nada adianta a conexão funcionar se o agente não conseguir usá-la.",
        [
            "Vincule a fonte PostgreSQL ao template “Atendimento Vendas”.",
            "Publique a nova versão.",
            "No Chat, pergunte algo que só possa ser respondido pelos dados — por "
            "exemplo, o preço de um produto que você sabe que está no banco.",
        ],
        [
            "A resposta traz o valor real que está no banco.",
            "Se você conferir na conversa, aparece que a ferramenta de consulta foi usada.",
            "O agente não inventa números.",
        ],
        "Cole o print da resposta e, se possível, do detalhe mostrando a consulta executada.",
    )

    r.bloco(
        "6.7",
        "Escrita no banco exige confirmação",
        "Escrever é diferente de ler. O agente não pode alterar dados sem o "
        "cliente autorizar explicitamente.",
        [
            "Em um template com escrita habilitada, peça para registrar algo — "
            "por exemplo, “registre uma venda de 2 unidades do produto X”.",
            "Quando o agente pedir confirmação, primeiro responda “não”.",
            "Depois repita o pedido e confirme com “sim, pode registrar”.",
        ],
        [
            "O agente pede confirmação antes de gravar.",
            "Ao responder “não”, nada é gravado — confira no banco.",
            "Ao confirmar, a gravação acontece e o agente avisa.",
        ],
        "Cole o print dos dois momentos: a recusa e a confirmação.",
    )

    r.bloco(
        "6.8",
        "Fontes de dados não vazam entre empresas",
        "Teste de segurança.",
        [
            "Entre na Empresa B e vá em Fontes de dados.",
        ],
        ["As fontes da Empresa A não aparecem."],
    )

    # =====================================================================
    r.h1("10. Suíte 7 — Memória do agente")
    r.p(
        "Memória é diferente de histórico. O histórico é o que foi dito nesta "
        "conversa. A memória é o que o agente guarda para as PRÓXIMAS conversas — "
        "por exemplo, que aquele cliente prefere entrega pela manhã. Para testar, "
        "você precisa fechar uma conversa e abrir outra."
    )

    r.bloco(
        "7.1",
        "Criar uma memória",
        "Sem isso o agente recomeça do zero toda vez.",
        [
            "Abra uma conversa nova.",
            "Diga: “Anote que eu prefiro receber os relatórios sempre às segundas-feiras”.",
            "Continue conversando por mais duas ou três mensagens e encerre a conversa.",
            "Vá em Memórias.",
        ],
        [
            "A informação aparece na lista de memórias — pode demorar alguns segundos.",
            "O texto guardado faz sentido e não é a conversa inteira copiada.",
        ],
        "Cole o print da tela de Memórias com o registro.",
    )

    r.bloco(
        "7.2",
        "O agente usa a memória em outra conversa",
        "É o que prova que a memória serve para alguma coisa.",
        [
            "Abra uma conversa NOVA (não continue a anterior).",
            "Pergunte: “Quando eu prefiro receber relatórios?”",
        ],
        [
            "O agente responde “às segundas-feiras”.",
            "Ele acerta mesmo sem você repetir a informação.",
        ],
        "Cole o print da nova conversa com a resposta correta.",
    )

    r.bloco(
        "7.3",
        "Apagar uma memória",
        "O cliente precisa poder corrigir o que o agente guardou errado.",
        [
            "Em Memórias, apague o registro criado.",
            "Abra outra conversa nova e faça a mesma pergunta.",
        ],
        [
            "A memória some da lista.",
            "O agente não sabe mais responder — o que é o comportamento certo.",
        ],
    )

    r.bloco(
        "7.4",
        "Memórias não vazam entre empresas",
        "Memória guarda informação sensível de cliente. Vazamento aqui é grave.",
        [
            "Entre na Empresa B e vá em Memórias.",
        ],
        ["Nenhuma memória da Empresa A aparece."],
    )

    # =====================================================================
    r.h1("11. Suíte 8 — Artefatos: gráfico, planilha e PDF")
    r.p(
        "Artefato é o arquivo que a conversa produz. A regra importante: o "
        "gráfico e a planilha são feitos a partir de uma consulta que já "
        "aconteceu — o agente não inventa os números na hora de desenhar."
    )

    r.bloco(
        "8.1",
        "Gerar um gráfico",
        "É o artefato mais pedido em demonstração.",
        [
            "Use o template “Analista de Dados”.",
            "Peça: “Me mostre um gráfico de vendas por produto”.",
        ],
        [
            "O gráfico aparece dentro da conversa.",
            "Os números batem com o que está no banco.",
            "É possível interagir com o gráfico (passar o mouse e ver valores).",
        ],
        "Cole o print do gráfico.",
    )

    r.bloco(
        "8.2",
        "Gerar uma planilha",
        "Cliente adora exportar para Excel.",
        [
            "Peça: “Exporte esses dados em planilha”.",
            "Baixe o arquivo e abra.",
        ],
        [
            "O download acontece.",
            "O arquivo abre no Excel sem aviso de corrompido.",
            "Os dados dentro batem com a conversa.",
        ],
        "Cole o print da planilha aberta.",
    )

    r.bloco(
        "8.3",
        "Gerar um PDF",
        "Fecha o conjunto de saídas.",
        ["Peça um relatório em PDF.", "Baixe e abra."],
        ["O PDF abre corretamente.", "O conteúdo corresponde ao que foi conversado."],
    )

    # =====================================================================
    r.h1("12. Suíte 9 — Pagamentos: Mercado Pago e PIX")
    r.p(
        "O agente consegue emitir uma cobrança PIX dentro da conversa. Esta "
        "suíte é sensível: envolve dinheiro. Use SEMPRE a credencial de sandbox "
        "(teste) enquanto estiver validando."
    )
    r.aviso(
        "Sandbox, não produção",
        "Ao cadastrar a credencial, escolha Ambiente = Sandbox (testes). Se você "
        "usar a credencial de produção, as cobranças geradas são reais e podem "
        "ser pagas de verdade. Confirme o ambiente antes de qualquer teste.",
        cor="FDECEA",
    )

    r.bloco(
        "9.1",
        "Cadastrar a credencial",
        "É o pré-requisito de tudo nesta suíte, e também um teste de segurança: "
        "o token não pode voltar para a tela.",
        [
            "Vá em Pagamentos.",
            "Preencha o access token do Mercado Pago e o segredo do webhook.",
            "Escolha Ambiente = Sandbox.",
            "Salve e recarregue a página.",
        ],
        [
            "A credencial é salva.",
            "Depois de recarregar, o token NÃO aparece na tela — nem inteiro, nem parcial.",
            "A tela indica que existe credencial configurada.",
        ],
        "Cole o print da tela após salvar e recarregar.",
    )

    r.bloco(
        "9.2",
        "O agente gera uma cobrança PIX",
        "É a funcionalidade em si.",
        [
            "Use um template que tenha as ferramentas de cobrança marcadas.",
            "Peça: “Gere uma cobrança PIX de R$ 12,34 para o pedido TESTE-QA”.",
        ],
        [
            "O agente gera a cobrança.",
            "Aparece o código PIX (copia e cola) ou o QR Code.",
            "O valor é exatamente R$ 12,34 — nem centavo a mais, nem a menos.",
            "A cobrança aparece na lista em Pagamentos.",
        ],
        "Cole o print da cobrança gerada na conversa e o print da lista em Pagamentos.",
    )

    r.aviso(
        "Teste do valor — leia com atenção",
        "Peça cobranças de valores diferentes (R$ 1,00, R$ 99,90, R$ 1.234,56) e "
        "confira cada uma. Um sistema parecido no mercado tinha o valor fixo em "
        "R$ 0,01 escondido no código: todas as cobranças saíam com um centavo. "
        "Esse é exatamente o tipo de erro que só o teste manual pega.",
    )

    r.bloco(
        "9.3",
        "Pagar a cobrança no sandbox",
        "Gerar é metade. O sistema precisa perceber o pagamento.",
        [
            "Pegue o código PIX gerado.",
            "Use o comprador de teste do Mercado Pago (sandbox) para pagar.",
            "Volte à conversa e pergunte: “O pagamento do pedido TESTE-QA já caiu?”",
        ],
        [
            "O agente consulta e responde que foi pago.",
            "O status na tela de Pagamentos muda para pago/aprovado.",
            "O valor confirmado é o mesmo que foi cobrado.",
        ],
        "Cole o print do pagamento no sandbox e o print do status atualizado.",
    )

    r.bloco(
        "9.4",
        "Consultar cobrança que não foi paga",
        "O caminho negativo também precisa funcionar.",
        [
            "Gere uma cobrança nova e NÃO pague.",
            "Pergunte ao agente se ela foi paga.",
        ],
        [
            "O agente responde que ainda está pendente.",
            "Ele não afirma que foi paga.",
        ],
    )

    r.bloco(
        "9.5",
        "Credenciais de pagamento não vazam entre empresas",
        "Teste de segurança — aqui o vazamento envolveria dinheiro de terceiro.",
        [
            "Entre na Empresa B e vá em Pagamentos.",
        ],
        [
            "A credencial da Empresa A não aparece.",
            "As cobranças da Empresa A não aparecem.",
        ],
    )

    # =====================================================================
    r.h1("13. Suíte 10 — Integrações e API pública")
    r.p(
        "A API pública permite que um sistema do cliente converse com o agente "
        "sem abrir o navegador. É por aqui que muita integração de cliente "
        "acontece, e é também o que a W-API usa."
    )

    r.bloco(
        "10.1",
        "Criar uma credencial de API",
        "A chave só aparece uma vez. Se o sistema mostrar de novo depois, é falha.",
        [
            "Vá em Integrações.",
            "Crie uma integração do tipo API, escolhendo um template padrão.",
            "Copie a credencial mostrada e guarde em local seguro.",
            "Recarregue a página.",
        ],
        [
            "A credencial aparece uma única vez, com aviso claro de que não será mostrada de novo.",
            "Depois de recarregar, ela não aparece mais.",
            "A integração continua listada.",
        ],
        "Cole o print do momento em que a credencial é exibida (pode ocultar parte dela no print).",
    )

    r.bloco(
        "10.2",
        "Usar a API para conversar",
        "Prova que a integração funciona de fora do navegador.",
        [
            "Peça ajuda ao responsável técnico para executar a chamada abaixo, "
            "trocando a credencial pela sua.",
            "Observe a resposta.",
        ],
        [
            "A resposta vem com o texto do agente.",
            "O comportamento é igual ao do chat na tela.",
        ],
    )
    r.codigo(
        f"curl -X POST {BACKEND}/v1/messages \\\n"
        '  -H "Authorization: Bearer SUA_CREDENCIAL" \\\n'
        '  -H "Content-Type: application/json" \\\n'
        '  -d \'{"message": "Olá, você está funcionando?", "mode": "sync"}\''
    )

    r.bloco(
        "10.3",
        "Credencial inválida é recusada",
        "Teste de segurança básico e obrigatório.",
        [
            "Repita a chamada acima trocando a credencial por “credencial-invalida”.",
        ],
        [
            "A chamada é recusada.",
            "A resposta não traz dado nenhum de nenhuma empresa.",
        ],
    )

    # =====================================================================
    r.h1("14. Suíte 11 — WhatsApp (W-API) em várias empresas")
    r.p(
        "Aqui começa a parte mais próxima do uso real. Cada empresa conecta o "
        "próprio número. O teste importante é ter DUAS empresas com números "
        "diferentes ao mesmo tempo, e provar que uma não recebe a mensagem da outra."
    )

    r.bloco(
        "11.1",
        "Conectar o WhatsApp da Empresa A",
        "É o caminho feliz. Serve de base para os testes seguintes.",
        [
            "Na Empresa A, vá em Integrações e crie uma integração do tipo WhatsApp (W-API).",
            "Abra a conexão e preencha Instance ID, Token da instância e URL base da W-API.",
            "Salve.",
            "Copie a URL de webhook mostrada na tela.",
            "No painel da W-API, cole essa URL no campo de webhook da instância.",
            "Conecte o número escaneando o QR Code pelo celular.",
        ],
        [
            "A conexão é salva.",
            "Ao reabrir, o token NÃO aparece — só a indicação de que já existe.",
            "A W-API mostra a instância como conectada.",
        ],
        "Cole o print da conexão salva e o print da W-API mostrando conectado.",
    )

    r.bloco(
        "11.2",
        "Receber e responder mensagem no WhatsApp",
        "É o teste mais importante desta suíte: a mensagem sai do celular e volta respondida.",
        [
            "Do seu celular pessoal (o número “cliente”), mande uma mensagem para o "
            "número conectado da Empresa A.",
            "Aguarde a resposta.",
        ],
        [
            "A resposta chega no WhatsApp em tempo razoável.",
            "O conteúdo corresponde ao template configurado.",
            "A conversa também aparece do lado da plataforma.",
        ],
        "Cole o print da conversa no celular.",
    )

    r.bloco(
        "11.3",
        "Conectar o WhatsApp da Empresa B com outro número",
        "Duas empresas ao mesmo tempo é o cenário real do negócio.",
        [
            "Repita o teste 11.1 na Empresa B, com uma segunda instância da W-API "
            "e um segundo número.",
        ],
        [
            "As duas conexões coexistem sem conflito.",
            "Cada empresa mostra apenas o próprio número.",
        ],
    )

    r.bloco(
        "11.4",
        "Isolamento entre os dois números",
        "Este é o teste que protege o negócio. Uma mensagem que aparece na "
        "empresa errada é o pior tipo de falha possível neste sistema.",
        [
            "Mande uma mensagem para o número da Empresa A.",
            "Mande outra, com texto claramente diferente, para o número da Empresa B.",
            "Abra as duas empresas na plataforma e compare.",
        ],
        [
            "A mensagem do número A aparece SOMENTE na Empresa A.",
            "A mensagem do número B aparece SOMENTE na Empresa B.",
            "As respostas voltam cada uma para o número certo.",
            "O agente que responde é o template da empresa correspondente.",
        ],
        "Cole prints das duas empresas, mostrando que cada uma só tem a sua mensagem.",
    )

    r.bloco(
        "11.5",
        "Mensagem com mídia",
        "Cliente manda foto de nota fiscal, áudio, documento. Nada pode derrubar o fluxo.",
        [
            "Envie para o número da Empresa A: uma foto, um áudio e um documento PDF.",
        ],
        [
            "Nenhum dos envios derruba o atendimento.",
            "O sistema responde algo coerente — mesmo que seja para dizer que não "
            "consegue processar aquele tipo de mídia.",
            "Não fica sem resposta nenhuma.",
        ],
        "Cole o print da conversa com as mídias enviadas.",
    )

    r.bloco(
        "11.6",
        "Mensagem enviada com o sistema fora do ar",
        "Serve para descobrir se a mensagem se perde ou é recuperada.",
        [
            "Peça ao responsável técnico para reiniciar o serviço (ou avise que vai testar).",
            "Durante a reinicialização, mande uma mensagem para o número.",
            "Aguarde o sistema voltar.",
        ],
        [
            "A mensagem é respondida quando o sistema volta, OU",
            "fica registrado em algum lugar que ela chegou.",
            "Anote exatamente o que aconteceu — este é um teste de descoberta, "
            "e o resultado é informação valiosa mesmo que não seja o ideal.",
        ],
    )

    # =====================================================================
    r.h1("15. Suíte 12 — Atendimento omnichannel (Chatwoot) com 3 empresas")
    r.p(
        "O Chatwoot é o painel onde os atendentes humanos trabalham. A IA "
        "responde primeiro; quando não dá conta, passa para uma pessoa. O "
        "desenho é este:"
    )
    r.codigo(
        "cliente no WhatsApp\n"
        "      │\n"
        "      ▼\n"
        "  Chatwoot (caixa de entrada da empresa)\n"
        "      │\n"
        "      ├── IA responde  ──────────────▶ cliente\n"
        "      │\n"
        "      └── IA pede ajuda / humano escreve\n"
        "                  │\n"
        "                  ▼\n"
        "          atendente humano assume  ──▶ cliente\n"
        "          (a IA se cala sozinha)"
    )

    r.bloco(
        "12.1",
        "Provisionar o atendimento da Empresa A",
        "Cada empresa ganha a própria operação dentro do Chatwoot.",
        [
            "Na Empresa A, vá em Atendimento.",
            "Clique em Criar operação de atendimento.",
            "Depois clique em Abrir atendimento.",
        ],
        [
            "O status muda para Provisionado.",
            "O Chatwoot abre já logado — sem pedir senha nova.",
            "Você entra como administrador da conta da Empresa A.",
        ],
        "Cole o print do Chatwoot aberto, mostrando o nome da empresa.",
    )

    r.bloco(
        "12.2",
        "Provisionar Empresas B e C",
        "Três operações simultâneas é o cenário que o dono pediu para validar.",
        ["Repita o teste 12.1 nas Empresas B e C."],
        [
            "As três operações existem.",
            "Cada uma abre em uma conta separada do Chatwoot.",
        ],
    )

    r.bloco(
        "12.3",
        "Isolamento entre as três operações",
        "É o teste central desta suíte.",
        [
            "Abra o atendimento da Empresa A e anote o que aparece.",
            "Em outra aba, abra o da Empresa B.",
            "Em outra aba, o da Empresa C.",
            "Compare as três.",
        ],
        [
            "Cada uma mostra apenas as próprias conversas.",
            "Nenhuma lista contatos, caixas de entrada ou agentes das outras.",
            "Trocar de aba não mistura os dados.",
        ],
        "Cole prints das três abas.",
    )

    r.bloco(
        "12.4",
        "Conversa atendida pela IA",
        "É o fluxo normal, que deve resolver a maioria dos casos sem humano.",
        [
            "Mande uma mensagem simples pelo WhatsApp da Empresa A.",
            "Acompanhe a conversa aparecendo no Chatwoot.",
        ],
        [
            "A conversa aparece no Chatwoot.",
            "A resposta da IA aparece na conversa, identificada como do bot.",
            "O cliente recebe a resposta no WhatsApp.",
        ],
        "Cole o print da conversa no Chatwoot.",
    )

    r.bloco(
        "12.5",
        "Passagem para atendente humano",
        "É o momento mais delicado do produto: a IA precisa sair de cena e não voltar.",
        [
            "Na mesma conversa, peça algo que a IA não consiga resolver — por exemplo, "
            "“quero falar com um atendente humano”.",
            "Observe o Chatwoot.",
            "Assuma a conversa como atendente e responda manualmente.",
            "Mande mais uma mensagem pelo WhatsApp.",
        ],
        [
            "A conversa é encaminhada para atendimento humano.",
            "Fica registrado na conversa que houve a passagem.",
            "Depois que o humano responde, a IA NÃO responde mais por cima.",
            "A mensagem seguinte do cliente vai para o humano, não para a IA.",
        ],
        "Cole o print mostrando a passagem e o print da resposta humana.",
    )

    r.bloco(
        "12.6",
        "Três empresas conversando ao mesmo tempo",
        "Simula o dia real. Problemas de mistura de dados aparecem sob simultaneidade.",
        [
            "Peça ajuda a dois colegas — ou use três celulares.",
            "Mande mensagens ao mesmo tempo para os números das Empresas A, B e C.",
            "Faça cada uma perguntar algo bem diferente e reconhecível.",
        ],
        [
            "Cada resposta chega no número certo.",
            "Nenhuma conversa aparece na empresa errada.",
            "As respostas correspondem ao template de cada empresa.",
            "Nada fica sem resposta.",
        ],
        "Cole prints dos três celulares e das três telas do Chatwoot.",
    )

    r.h2("15.1 Fluxo sem IA — menu de opções (“digite 1, digite 2”)")
    r.aviso(
        "Leia antes de testar: isto ainda não existe pronto",
        "Um menu numérico determinístico (“digite 1 para vendas, 2 para suporte”) "
        "NÃO é uma funcionalidade instalada hoje. O Chatwoot da versão em uso "
        "(3.16) não traz um construtor de menus, e a ponte com a plataforma "
        "sempre encaminha a mensagem para a IA ou para um humano. Os testes "
        "abaixo servem para mapear o que dá para fazer hoje e registrar a "
        "diferença para o que o cliente pode pedir. Não marque como bug o que "
        "simplesmente não foi construído — marque como “não suportado”.",
    )

    r.bloco(
        "12.7",
        "Menu simulado por prompt (o que dá para fazer hoje)",
        "É a alternativa disponível: a IA apresenta o menu e interpreta a "
        "resposta. Funciona bem na prática, mas depende do modelo — não é "
        "determinístico.",
        [
            "Crie um template chamado “Menu Recepção”.",
            "No prompt do supervisor, escreva algo como: “Sempre inicie a conversa "
            "oferecendo exatamente estas opções: 1 - Vendas, 2 - Suporte, 3 - "
            "Financeiro. Se a pessoa responder 1, 2 ou 3, siga o assunto "
            "correspondente. Se responder outra coisa, repita o menu uma vez.”",
            "Publique e vincule ao WhatsApp de uma das empresas.",
            "Mande uma mensagem e responda “2”.",
            "Repita mandando “banana” em vez de um número.",
        ],
        [
            "A primeira resposta traz o menu com as três opções.",
            "Responder “2” leva ao assunto de suporte.",
            "Responder algo fora do menu faz o agente repetir as opções.",
            "Anote se em alguma tentativa o agente ignorou o menu — isso é a "
            "limitação esperada de um menu feito por IA, e é informação importante para o produto.",
        ],
        "Cole o print da conversa com o menu e das duas respostas testadas.",
    )

    r.bloco(
        "12.8",
        "Repetir o menu 10 vezes e medir a consistência",
        "Como o menu depende do modelo, o que interessa é a taxa de acerto. Este "
        "número é o que vai embasar a resposta ao cliente que pedir “menuzinho”.",
        [
            "Abra 10 conversas novas (pode apagar o contato e mandar de novo).",
            "Em todas, responda “1”.",
            "Anote quantas vezes o agente seguiu corretamente para Vendas.",
        ],
        [
            "Registre o número: ____ de 10 acertos.",
            "Anote qualquer resposta estranha, copiando o texto exato.",
        ],
        "Cole prints das tentativas que deram errado, se houver.",
    )

    r.bloco(
        "12.9",
        "Atendimento só humano, sem IA nenhuma",
        "Alguns clientes vão querer o Chatwoot puro, sem agente. Vale saber se dá.",
        [
            "Peça ao responsável técnico para desligar o piloto automático da IA "
            "em uma das empresas.",
            "Mande uma mensagem pelo WhatsApp dessa empresa.",
        ],
        [
            "A mensagem chega ao Chatwoot.",
            "A IA NÃO responde.",
            "Um atendente humano consegue responder normalmente.",
        ],
    )

    # =====================================================================
    r.h1("16. Suíte 13 — Outras telas")

    r.bloco(
        "13.1",
        "Arquivos e busca em documentos",
        "Permite que o agente responda com base em documentos da empresa.",
        [
            "Vá em Arquivos e suba um PDF com informação específica e verificável.",
            "Vincule o arquivo a um agente do template.",
            "Publique e pergunte, no Chat, algo que só esteja naquele documento.",
        ],
        [
            "O upload conclui.",
            "O agente responde com a informação correta do documento.",
            "Perguntando algo que não está no documento, ele não inventa.",
        ],
        "Cole o print da resposta baseada no documento.",
    )

    r.bloco(
        "13.2",
        "Consumo",
        "Mostra quanto a empresa está gastando.",
        ["Vá em Consumo depois de ter feito várias conversas."],
        [
            "Aparecem números de uso.",
            "Os números aumentam depois de novas conversas.",
            "Os dados são só da empresa logada.",
        ],
        "Cole o print da tela de Consumo.",
    )

    r.bloco(
        "13.3",
        "Personalizar (marca da empresa)",
        "É o que faz o cliente sentir que o sistema é dele.",
        [
            "Vá em Personalizar e troque cor e logo.",
            "Salve e recarregue.",
        ],
        [
            "As mudanças aparecem no topo da tela.",
            "Elas continuam depois de sair e entrar de novo.",
            "A Empresa B continua com a aparência dela — a mudança não vazou.",
        ],
        "Cole prints antes e depois, e o print da Empresa B inalterada.",
    )

    r.bloco(
        "13.4",
        "MCP Store",
        "Permite ligar ferramentas externas ao agente.",
        ["Vá em MCP Store e veja o catálogo.", "Ative um servidor, se houver algum disponível."],
        [
            "O catálogo carrega.",
            "A ativação aparece como ativa.",
            "As ferramentas do servidor aparecem no editor de template.",
        ],
    )

    # =====================================================================
    r.h1("17. Suíte 14 — Segurança e isolamento (a mais importante)")
    r.p(
        "Se todas as outras suítes passarem e esta falhar, o sistema não pode ir "
        "para produção. Aqui você tenta, de propósito, ver o dado de uma empresa "
        "estando logado em outra. Faça com atenção: um resultado positivo aqui "
        "(ou seja, você conseguiu ver) é o achado mais valioso de todo o roteiro."
    )
    r.aviso(
        "Como testar “trocando o endereço”",
        "Vários testes pedem que você troque um identificador na barra de "
        "endereço. Para pegar um identificador da outra empresa, entre nela, "
        "copie o trecho da URL, saia, entre na primeira empresa e cole. Se o "
        "sistema mostrar o dado, é falha grave — registre imediatamente e avise "
        "o responsável no mesmo dia.",
        cor="FDECEA",
    )

    r.tabela(
        ["Tela", "O que tentar ver da outra empresa", "Resultado esperado"],
        [
            [
                "Templates",
                "Abrir um template da outra empresa pela URL",
                "Bloqueado ou não encontrado",
            ],
            [
                "Fontes de dados",
                "Abrir uma fonte da outra empresa pela URL",
                "Bloqueado ou não encontrado",
            ],
            ["Chat", "Abrir uma conversa da outra empresa pela URL", "Bloqueado ou não encontrado"],
            [
                "Memórias",
                "Listar memórias estando na outra empresa",
                "Lista vazia ou só as próprias",
            ],
            ["Pagamentos", "Ver credencial e cobranças da outra", "Não aparecem"],
            ["Serviços de IA", "Ver contas e combos da outra", "Não aparecem"],
            ["Usuários", "Listar usuários da outra empresa", "Não aparecem"],
            ["Arquivos", "Baixar um arquivo da outra empresa pela URL", "Bloqueado"],
            ["Atendimento", "Abrir o Chatwoot da outra empresa", "Bloqueado"],
        ],
        larguras=[3.6, 7.4, 6.0],
    )

    r.bloco(
        "14.1",
        "Varredura completa de isolamento",
        "Percorrer a tabela acima item por item. É repetitivo de propósito.",
        [
            "Entre na Empresa A e anote os identificadores das URLs de cada tela da tabela.",
            "Saia e entre na Empresa B.",
            "Tente abrir cada um dos endereços anotados.",
            "Marque na tabela acima o resultado de cada um.",
        ],
        [
            "TODOS os itens da tabela resultam em bloqueio ou “não encontrado”.",
            "Nenhum dado da Empresa A aparece.",
            "Nenhuma tela mostra erro técnico cru.",
        ],
        "Cole prints de pelo menos três tentativas bloqueadas.",
    )

    r.bloco(
        "14.2",
        "Segredos nunca voltam para a tela",
        "Uma vez salvo, um segredo não deve ser exibido de novo por nada.",
        [
            "Percorra: senha de fonte de dados, token do Mercado Pago, token da W-API, "
            "chave de API de IA, credencial de integração.",
            "Em cada uma, salve, recarregue a página e observe.",
            "Abra também o menu do navegador (F12 > aba Rede) e veja a resposta da tela.",
        ],
        [
            "Nenhum segredo aparece na tela.",
            "Nenhum segredo aparece na resposta técnica vista no F12.",
            "O sistema indica apenas que existe um valor salvo.",
        ],
        "Cole o print do F12 mostrando a resposta sem o segredo.",
    )

    r.bloco(
        "14.3",
        "Usuário removido perde o acesso na hora",
        "Desligamento de funcionário é um caso real e frequente.",
        [
            "Crie um usuário na Empresa A e entre com ele em outro navegador.",
            "No primeiro navegador, como administrador, remova esse usuário.",
            "No segundo navegador, tente navegar ou recarregar.",
        ],
        [
            "O acesso é encerrado.",
            "O sistema pede login de novo.",
            "Ele não consegue mais entrar.",
        ],
    )

    # =====================================================================
    r.h1("18. Como registrar um problema")
    r.p(
        "Um bug bem descrito é resolvido em minutos; um bug mal descrito volta "
        "três vezes para pedir informação. Use o modelo abaixo — copie e "
        "preencha um para cada problema encontrado."
    )
    r.tabela(
        ["Campo", "Preencher"],
        [
            ["Código do teste", "Ex.: 11.4"],
            ["Título curto", "Uma frase objetiva do que está errado"],
            ["Empresa / usuário usado", ""],
            ["O que eu fiz (passo a passo)", "Detalhado o suficiente para outra pessoa repetir"],
            ["O que eu esperava", ""],
            ["O que aconteceu de verdade", ""],
            ["Acontece sempre?", "( ) Sempre  ( ) Às vezes  ( ) Só uma vez"],
            ["Data e hora exata", "Ajuda a achar o registro técnico"],
            ["Prints", "Anexar"],
            ["Gravidade", "( ) Impede o uso  ( ) Atrapalha  ( ) Detalhe"],
        ],
        larguras=[5.0, 12.0],
    )
    r.p("Regras para classificar a gravidade:", negrito=True)
    r.bullet(
        "Impede o uso: dado de uma empresa aparece em outra; segredo exposto; "
        "não dá para logar; cobrança com valor errado; mensagem entregue à empresa errada."
    )
    r.bullet("Atrapalha: funciona, mas com erro visível, lentidão grande ou passo confuso.")
    r.bullet("Detalhe: texto errado, alinhamento, cor, algo que não impede o trabalho.")

    r.aviso(
        "Achou vazamento entre empresas? Avise no mesmo dia",
        "Qualquer caso em que uma empresa enxergue dado de outra deve ser "
        "comunicado imediatamente ao responsável técnico, sem esperar o fim do "
        "roteiro. Esse é o único tipo de problema que justifica interromper os testes.",
        cor="FDECEA",
    )

    # =====================================================================
    r.h1("19. Fechamento")
    r.h2("19.1 Resumo dos resultados")
    r.tabela(
        ["Suíte", "Testes", "Passou", "Falhou", "Bloqueado", "Observações"],
        [
            ["1 — Acesso e permissões", "5", "", "", "", ""],
            ["2 — Empresas", "6", "", "", "", ""],
            ["3 — Serviços de IA", "8", "", "", "", ""],
            ["4 — Templates", "5", "", "", "", ""],
            ["5 — Chat", "5", "", "", "", ""],
            ["6 — Fontes de dados", "8", "", "", "", ""],
            ["7 — Memória", "4", "", "", "", ""],
            ["8 — Artefatos", "3", "", "", "", ""],
            ["9 — Pagamentos", "5", "", "", "", ""],
            ["10 — Integrações e API", "3", "", "", "", ""],
            ["11 — WhatsApp", "6", "", "", "", ""],
            ["12 — Atendimento (Chatwoot)", "9", "", "", "", ""],
            ["13 — Outras telas", "4", "", "", "", ""],
            ["14 — Segurança", "3", "", "", "", ""],
            ["TOTAL", "74", "", "", "", ""],
        ],
        larguras=[5.0, 1.8, 1.8, 1.8, 2.0, 4.6],
    )

    r.h2("19.2 Parecer final")
    r.p(
        "Escreva com suas palavras: o sistema está pronto para ser usado por um "
        "cliente real? O que mais te preocupou? O que funcionou melhor do que "
        "você esperava? Esta parte é lida com atenção — a impressão de quem "
        "testou pela primeira vez é justamente o que a equipe que construiu já "
        "não consegue enxergar."
    )
    for _ in range(10):
        r.doc.add_paragraph("_" * 95)

    r.h2("19.3 O que estudar depois deste roteiro")
    r.p(
        "Você vai cuidar deste sistema e apresentá-lo a clientes. Depois de "
        "rodar o roteiro, estes são os assuntos que mais rendem:"
    )
    r.bullet(
        "Como um agente decide usar uma ferramenta — entender isso explica 90% "
        "das perguntas de cliente sobre “por que ele não fez o que pedi”."
    )
    r.bullet("A diferença entre memória e histórico — é a dúvida mais comum em demonstração.")
    r.bullet(
        "Por que cada empresa tem instância de IA separada — é o argumento de "
        "venda mais forte do produto, e você precisa saber explicá-lo em uma frase."
    )
    r.bullet(
        "O caminho de uma mensagem de WhatsApp até a resposta — desenhe esse "
        "fluxo de memória; é o que mais te ajuda a diagnosticar problema em campo."
    )
    r.bullet(
        "Quando a IA deve sair de cena e chamar um humano — é o que separa um "
        "atendimento bom de um cliente irritado."
    )

    r.doc.add_paragraph()
    r.tabela(
        ["Responsável pelos testes", "Data", "Assinatura"],
        [["", "", ""], ["", "", ""]],
        larguras=[7.0, 4.0, 6.0],
    )

    return r


def main() -> None:
    destino = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("docs/Roteiro-QA-Plataforma.docx")
    destino.parent.mkdir(parents=True, exist_ok=True)
    montar().salvar(destino)
    print(f"gerado: {destino.resolve()}")


if __name__ == "__main__":
    main()
